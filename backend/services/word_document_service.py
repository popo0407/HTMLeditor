"""
Wordファイル生成サービス

責務: テキストデータからWordファイル (.docx) を生成
"""

from docx import Document
from docx.shared import Inches
from io import BytesIO
import logging
from typing import Dict, Any

logger = logging.getLogger(__name__)


class WordDocumentService:
    """Wordドキュメント生成サービス"""

    @staticmethod
    def create_document_from_text(text_content: str, meeting_info: Dict[str, Any] = None) -> bytes:
        """
        テキストからWordドキュメントを生成
        
        Args:
            text_content: テキスト内容
            meeting_info: 会議情報（ヘッダー用）
            
        Returns:
            Wordファイルのバイトデータ
        """
        try:
            # 新しいドキュメントを作成
            doc = Document()
            
            # 会議情報がある場合はヘッダーを追加
            if meeting_info:
                # タイトル
                title = meeting_info.get('会議タイトル') or meeting_info.get('title') or '元データ'
                title_paragraph = doc.add_heading(title, level=1)
                
                # 会議情報テーブル
                if any([
                    meeting_info.get('会議日時') or meeting_info.get('datetime'),
                    meeting_info.get('会議場所') or meeting_info.get('location'),
                    meeting_info.get('参加者') or meeting_info.get('participants')
                ]):
                    info_table = doc.add_table(rows=0, cols=2)
                    info_table.style = 'Table Grid'
                    
                    # 会議日時
                    datetime_str = meeting_info.get('会議日時') or meeting_info.get('datetime')
                    if datetime_str:
                        row = info_table.add_row()
                        row.cells[0].text = '会議日時'
                        row.cells[1].text = str(datetime_str)
                    
                    # 会議場所
                    location = meeting_info.get('会議場所') or meeting_info.get('location')
                    if location:
                        row = info_table.add_row()
                        row.cells[0].text = '会議場所'
                        row.cells[1].text = str(location)
                    
                    # 参加者
                    participants = meeting_info.get('参加者') or meeting_info.get('participants')
                    if participants:
                        row = info_table.add_row()
                        row.cells[0].text = '参加者'
                        if isinstance(participants, list):
                            row.cells[1].text = ', '.join(participants)
                        else:
                            row.cells[1].text = str(participants)
                
                # 区切り線
                doc.add_paragraph('─' * 50)
            
            # メインコンテンツ
            content_heading = doc.add_heading('内容', level=2)
            
            # テキスト内容を段落として追加
            if text_content:
                # BOMを除去
                clean_text = text_content.lstrip('\ufeff')
                
                # 改行で分割して段落を作成
                lines = clean_text.split('\n')
                for line in lines:
                    if line.strip():  # 空行でない場合のみ追加
                        doc.add_paragraph(line.strip())
                    else:
                        doc.add_paragraph()  # 空行を追加
            else:
                doc.add_paragraph('（内容なし）')
            
            # バイトストリームに保存
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            
            logger.info(f"Wordドキュメント生成完了: {len(bio.getvalue())} bytes")
            
            return bio.getvalue()
            
        except Exception as e:
            logger.error(f"Wordドキュメント生成エラー: {e}")
            raise Exception(f"Wordドキュメント生成に失敗しました: {e}")

    @staticmethod
    def create_meeting_minutes_document(meeting_info: Dict[str, Any], minutes_html: str) -> bytes:
        """
        会議情報とHTML議事録からWordドキュメントを生成
        
        Args:
            meeting_info: 会議情報
            minutes_html: 議事録HTML
            
        Returns:
            Wordファイルのバイトデータ
        """
        try:
            from bs4 import BeautifulSoup
            
            # HTMLからテキストを抽出
            soup = BeautifulSoup(minutes_html, 'html.parser')
            text_content = soup.get_text(separator='\n', strip=True)
            
            return WordDocumentService.create_document_from_text(text_content, meeting_info)
            
        except ImportError:
            # BeautifulSoupがない場合はHTMLタグを簡易除去
            import re
            text_content = re.sub(r'<[^>]+>', '', minutes_html)
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            return WordDocumentService.create_document_from_text(text_content, meeting_info)
        except Exception as e:
            logger.error(f"議事録Wordドキュメント生成エラー: {e}")
            raise Exception(f"議事録Wordドキュメント生成に失敗しました: {e}")
