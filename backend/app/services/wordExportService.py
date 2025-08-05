"""
Wordファイル出力サービス

開発憲章の「関心の分離」に従い、
Wordファイル出力を独立したサービスで管理
"""

import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup
import re


class WordExportService:
    """
    HTMLコンテンツをWordファイルに変換するサービス
    """
    
    @staticmethod
    def html_to_word(html_content: str, title: str = "エクスポートされたドキュメント") -> bytes:
        """
        HTMLコンテンツをWordファイルに変換
        
        Args:
            html_content: HTMLコンテンツ
            title: ドキュメントタイトル
            
        Returns:
            Wordファイルのバイトデータ
        """
        try:
            # 新しいWordドキュメントを作成
            doc = Document()
            
            # ドキュメントのプロパティを設定
            doc.core_properties.title = title
            doc.core_properties.creator = "HTMLエディタ"
            doc.core_properties.description = "TinyMCEエディタからエクスポート"
            
            # HTMLをパース
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # スタイルタグを除去（Wordでは使用しない）
            for style in soup.find_all('style'):
                style.decompose()
            
            # スクリプトタグを除去
            for script in soup.find_all('script'):
                script.decompose()
            
            # 各要素を処理
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'blockquote']):
                WordExportService._process_element(doc, element)
            
            # ファイルをバイトストリームに出力
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return output.getvalue()
            
        except Exception as e:
            raise Exception(f"Wordファイルの作成に失敗しました: {str(e)}")
    
    @staticmethod
    def _process_element(doc: Document, element):
        """
        個別のHTML要素をWordドキュメントに追加
        
        Args:
            doc: Wordドキュメント
            element: HTML要素
        """
        tag_name = element.name
        
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # 見出し
            heading = doc.add_heading(element.get_text(), level=int(tag_name[1]))
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        elif tag_name == 'p':
            # 段落
            paragraph = doc.add_paragraph()
            WordExportService._process_inline_elements(paragraph, element)
            
        elif tag_name in ['ul', 'ol']:
            # リスト
            for li in element.find_all('li', recursive=False):
                list_item = doc.add_paragraph()
                list_item.style = 'List Bullet' if tag_name == 'ul' else 'List Number'
                WordExportService._process_inline_elements(list_item, li)
                
        elif tag_name == 'table':
            # テーブル
            WordExportService._process_table(doc, element)
            
        elif tag_name == 'blockquote':
            # 引用
            quote = doc.add_paragraph()
            quote.style = 'Quote'
            WordExportService._process_inline_elements(quote, element)
    
    @staticmethod
    def _process_inline_elements(paragraph, element):
        """
        インライン要素（太字、斜体等）を処理
        
        Args:
            paragraph: Word段落
            element: HTML要素
        """
        for content in element.contents:
            if hasattr(content, 'name'):
                # HTML要素の場合
                if content.name == 'strong' or content.name == 'b':
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name == 'em' or content.name == 'i':
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                elif content.name == 'code':
                    run = paragraph.add_run(content.get_text())
                    run.font.name = 'Courier New'
                else:
                    # その他の要素は再帰的に処理
                    WordExportService._process_inline_elements(paragraph, content)
            else:
                # テキストの場合
                if str(content).strip():
                    paragraph.add_run(str(content))
    
    @staticmethod
    def _process_table(doc: Document, table_element):
        """
        テーブルを処理
        
        Args:
            doc: Wordドキュメント
            table_element: テーブル要素
        """
        # 行数をカウント
        rows = table_element.find_all('tr')
        if not rows:
            return
            
        # 列数をカウント（最初の行から）
        first_row = rows[0]
        cols = first_row.find_all(['td', 'th'])
        if not cols:
            return
            
        # テーブルを作成
        table = doc.add_table(rows=len(rows), cols=len(cols))
        table.style = 'Table Grid'
        
        # 各セルを処理
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if i < len(table.rows) and j < len(table.rows[i].cells):
                    word_cell = table.rows[i].cells[j]
                    WordExportService._process_inline_elements(word_cell.paragraphs[0], cell)
                    
                    # ヘッダーセルの場合は太字にする
                    if cell.name == 'th':
                        word_cell.paragraphs[0].runs[0].bold = True 